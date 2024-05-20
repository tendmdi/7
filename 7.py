from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.widget import Widget
from kivy.graphics import Ellipse, Color
from kivy.uix.gridlayout import GridLayout
from module import mass
from module import plankla
from module import svet
from docx import Document
from openpyxl import Workbook
from abc import ABC, abstractmethod

def managed_property(func):
    return property(**func())

class Particle(ABC):
    def __init__(self, charge, mass):
        self._charge = charge
        self._mass = mass

    @managed_property
    def charge():
        def fget(self):
            return self._charge
        def fset(self, value):
            self._charge = value
        return locals()

    @managed_property
    def mass():
        def fget(self):
            return self._mass
        def fset(self, value):
            self._mass = value
        return locals()

    @abstractmethod
    def specific_charge(self):
        pass

    def __str__(self):
        return f"{self.__class__.__name__} with charge {self._charge} and mass {self._mass}"

    def __repr__(self):
        return f"{self.__class__.__name__}(charge={self._charge}, mass={self._mass})"

class Electron(Particle):
    def __init__(self):
        super().__init__(mass.ze, mass.me)

    def specific_charge(self):
        return self.charge / self.mass

class Proton(Particle):
    def __init__(self):
        super().__init__(mass.zp, mass.mp)

    def specific_charge(self):
        return self.charge / self.mass

class Neutron(Particle):
    def __init__(self):
        super().__init__(0, mass.mn)

    def specific_charge(self):
        return 0

selected_particle = None

def fu(particle):
    global selected_particle
    selected_particle = particle

class GraphicsWidget(Widget):
    def __init__(self, **kwargs):
        super(GraphicsWidget, self).__init__(**kwargs)
        self.bind(size=self._update_canvas, pos=self._update_canvas)
        self.draw_particles()

    def draw_particles(self):
        r = 50
        y_offset = 710
        with self.canvas:
            # электрон
            Color(1, 0, 0)
            Ellipse(pos=(200 - r, y_offset - r), size=(2 * r, 2 * r))
            
            # протон
            Color(0, 0, 1)
            Ellipse(pos=(800 - r, y_offset - r), size=(2 * r, 2 * r),)
            
            # нейтрон
            Color(0, 1, 0)
            Ellipse(pos=(500 - r, y_offset - r), size=(2 * r, 2 * r))

    def _update_canvas(self, *args):
        self.canvas.clear()
        self.draw_particles()

class MainLayout(BoxLayout):
    def __init__(self, **kwargs):
        super(MainLayout, self).__init__(orientation='vertical', **kwargs)

        self.graphics_widget = GraphicsWidget(size_hint=(1, 0.5))
        self.add_widget(self.graphics_widget)

        self.add_widget(Label(text="Выберите частицу:"))
        particle_layout = GridLayout(cols=3, size_hint_y=None, height=30)
        self.add_widget(particle_layout)
        particles = ['Протон', 'Электрон', 'Нейтрон']
        for particle in particles:
            particle_button = Button(text=particle, on_press=lambda instance, p=particle: fu(p))
            particle_layout.add_widget(particle_button)

        self.add_widget(Label(text="Заряд частицы:"))
        self.z_entry = TextInput(multiline=False)
        self.add_widget(self.z_entry)

        self.add_widget(Label(text="Масса частицы:"))
        self.m_entry = TextInput(multiline=False)
        self.add_widget(self.m_entry)

        self.shetik_button = Button(text="Рассчитать", on_press=self.calculate_charge)
        self.add_widget(self.shetik_button)

        self.result_label = Label(text="")
        self.add_widget(self.result_label)

    def calculate_charge(self, instance):
        global selected_particle
        if selected_particle == 'Протон':
            particle = Proton()
        elif selected_particle == 'Электрон':
            particle = Electron()
        elif selected_particle == 'Нейтрон':
            particle = Neutron()
        else:
            self.result_label.text = "Выберите частицу"
            return

        k = particle.specific_charge()
        self.result_label.text = f"Удельный заряд: {k}"

        volna = (plankla.h) / (svet.c * particle.mass)

        doc = Document()
        doc.add_paragraph(f'Удельный заряд частицы "{selected_particle}": {k}')
        doc.add_paragraph(f'Комптоновская длина волны частицы "{selected_particle}": {volna}')
        doc.save('вордик.docx')

        wb = Workbook()
        ws = wb.active
        ws.append(['Удельный заряд', k])
        ws.append(['Комптоновская длина волны', volna])
        wb.save('экселька.xlsx')

class MyApp(App):
    def build(self):
        return MainLayout()

if __name__ == "__main__":
    MyApp().run()
